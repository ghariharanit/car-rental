#!/usr/bin/env python3
"""Generate DriveEase project report as a Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "driveease-project-report.docx"
SCREENSHOTS_DIR = ROOT / "docs" / "screenshots"

FONT = "Times New Roman"
SECTION_SIZE = Pt(24)
BODY_SIZE = Pt(12)
HEADING_SIZE = Pt(14)


def set_run_font(run, size=BODY_SIZE, bold=False):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_section_heading(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, size=SECTION_SIZE, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_subheading(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, size=HEADING_SIZE, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, bold=True)
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run)
    doc.add_paragraph()
    return table


def add_figure(doc, filename, caption):
    """Embed a screenshot if present; otherwise leave a placeholder caption."""
    add_body(doc, caption)
    image_path = SCREENSHOTS_DIR / filename
    if image_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(5.75))
    else:
        add_body(doc, f"[Insert screenshot: docs/screenshots/{filename}]")
    doc.add_paragraph()


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE

    # Title page
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DriveEase — Car Rental Web Application")
    set_run_font(run, size=SECTION_SIZE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "A Next.js Full-Stack Application with Supabase Auth and PostgreSQL Backend"
    )
    set_run_font(run, size=HEADING_SIZE)

    doc.add_paragraph()
    for line in [
        "Submitted by: [Student Name]",
        "Registration No.: [Registration Number]",
        "Department: [Department Name]",
        "Institution: [College Name]",
        "Academic Year: [Academic Year]",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_run_font(run)

    add_page_break(doc)

    # DECLARATION
    add_section_heading(doc, "DECLARATION")
    add_body(
        doc,
        'I, [Student Name], hereby declare that the project work entitled '
        '"DriveEase — Car Rental Web Application" submitted to [College Name] '
        "is a record of an original work done by me under the guidance of "
        "[Guide Name], and this work has not been submitted elsewhere for "
        "the award of any degree or diploma.",
    )
    add_body(doc, "Place: [Place]")
    add_body(doc, "Date: [Date]")
    add_body(doc, "Signature: ___________________")
    add_page_break(doc)

    # ACKNOWLEDGEMENT
    add_section_heading(doc, "ACKNOWLEDGEMENT")
    add_body(
        doc,
        "I would like to express my sincere gratitude to my project guide, "
        "[Guide Name], for their valuable guidance, encouragement, and support "
        "throughout the development of this project.",
    )
    add_body(
        doc,
        "I am thankful to [College Name] and the faculty of [Department Name] "
        "for providing the resources and environment needed to complete this work.",
    )
    add_body(
        doc,
        "I also extend my thanks to my family and friends for their constant "
        "motivation during this project.",
    )
    add_page_break(doc)

    # INDEX
    add_section_heading(doc, "INDEX")
    add_table(
        doc,
        ["S.No", "Section", "Page"],
        [
            ("1", "Declaration", ""),
            ("2", "Acknowledgement", ""),
            ("3", "Index", ""),
            ("4", "Abstract", ""),
            ("5", "Introduction", ""),
            ("5.1", "Background", ""),
            ("5.2", "Problem Statement", ""),
            ("5.3", "Objectives", ""),
            ("5.4", "Scope of the Project", ""),
            ("5.5", "Key Features and System Capabilities", ""),
            ("6", "System Analysis", ""),
            ("6.1", "Existing System", ""),
            ("6.2", "Proposed System", ""),
            ("6.3", "User Roles", ""),
            ("6.4", "Feasibility Study", ""),
            ("7", "System Specification", ""),
            ("8", "Software Description", ""),
            ("9", "Project Description", ""),
            ("10", "System Testing", ""),
            ("11", "System Implementation", ""),
            ("11.6", "Network Privacy (Browser vs Server)", ""),
            ("11.7", "Source Code Repository", ""),
            ("12", "Conclusion and Future Enhancement", ""),
            ("13", "Appendix", ""),
            ("13.1", "Appendix G — Application Screenshots", ""),
            ("13.2", "Appendix J — Source Code Listing", ""),
            ("14", "Bibliography and References", ""),
        ],
    )
    add_page_break(doc)

    # ABSTRACT
    add_section_heading(doc, "ABSTRACT")
    add_body(
        doc,
        "DriveEase is a web-based car rental application built with Next.js 14 (App Router) "
        "and Supabase as the backend platform. The system allows visitors to browse a catalog "
        "of rental cars, apply filters by name, fuel type, seats, and price, and view detailed "
        "car information. Registered users authenticate through Supabase Auth, book cars by "
        "selecting pickup and return dates, and view their booking history. Administrators "
        "access a dedicated panel to view dashboard statistics, manage the car fleet, and "
        "update booking statuses.",
    )
    add_body(
        doc,
        "The application uses TypeScript for type safety, Tailwind CSS and shadcn/ui for "
        "responsive UI components, and Supabase PostgreSQL for persistent storage. "
        "Authentication is implemented using Supabase Auth with @supabase/ssr for "
        "server-side session management and middleware-based route protection. Row Level "
        "Security (RLS) policies enforce user and admin access at the database layer.",
    )
    add_body(
        doc,
        "All data fetching and mutations run on the Next.js server via Server Components "
        "and Server Actions. The client browser communicates only with the Next.js "
        "application and does not make direct network requests to Supabase.",
    )
    add_body(
        doc,
        "The application supports local development, automated testing with Vitest and "
        "Playwright, and cloud deployment to Vercel with environment-based Supabase "
        "configuration. DriveEase is designed as a production-style full-stack web "
        "application suitable for academic evaluation, portfolio presentation, and live "
        "demonstration on a cloud-hosted URL.",
    )
    add_page_break(doc)

    # INTRODUCTION
    add_section_heading(doc, "INTRODUCTION")
    add_subheading(doc, "5.1 Background")
    add_body(
        doc,
        "The car rental industry increasingly relies on digital platforms to manage fleet "
        "inventory, accept online bookings, and provide customers with self-service access "
        "to vehicle information and reservations. Web applications offer centralized "
        "catalogs, role-based access for staff and customers, and automated price "
        "calculation based on rental duration.",
    )
    add_body(
        doc,
        "Modern full-stack frameworks such as Next.js enable rapid development of such "
        "applications with server-side rendering, Server Actions, and integrated "
        "authentication backed by Supabase PostgreSQL and Row Level Security.",
    )

    add_subheading(doc, "5.2 Problem Statement")
    add_body(
        doc,
        "Traditional car rental operations often depend on manual phone bookings, paper "
        "records, and disconnected systems for fleet management. Customers lack a unified "
        "interface to browse available vehicles, compare prices, and track their "
        "reservations. Administrators need a single dashboard to monitor bookings and "
        "maintain the car catalog.",
    )
    add_body(
        doc,
        "There is a need for a structured web application that demonstrates the core flows "
        "of an online car rental service: catalog discovery, secure user authentication, "
        "database-backed booking management, and administrative control — backed by "
        "Supabase PostgreSQL and deployable to a cloud hosting platform such as Vercel.",
    )

    add_subheading(doc, "5.3 Objectives")
    objectives = [
        "Application shell: responsive navigation, footer, home hero, and mobile-friendly layout.",
        "Vehicle catalog: search and filter capabilities, detail pages with image galleries, data in Supabase.",
        "Authentication: Supabase Auth login with user and admin roles, session cookies, and protected routes.",
        "Booking workflow: date selection, automatic price calculation, and personal booking history in Supabase.",
        "Administration: dashboard statistics, cars CRUD, and booking status management.",
        "User experience: loading skeletons, custom 404, toast notifications, and empty states.",
        "Cloud-ready backend: Supabase PostgreSQL with RLS, server-side data access, schema setup, and seed script.",
    ]
    for obj in objectives:
        add_bullet(doc, obj)

    add_subheading(doc, "5.5 Key Features and System Capabilities")
    add_table(
        doc,
        ["Module", "Capability", "Technology"],
        [
            ("Layout", "Navbar, footer, mobile nav, home hero", "Next.js, Tailwind, shadcn/ui"),
            ("Catalog", "Ten-vehicle grid, search/filter, detail galleries", "Supabase cars table"),
            ("Authentication", "Email/password login, role redirect, session cookies", "Supabase Auth, @supabase/ssr"),
            ("Booking", "Date selection, price calculation, booking history", "Supabase bookings table"),
            ("Administration", "Dashboard, cars CRUD, booking status", "RLS admin policies"),
            ("Backend", "PostgreSQL, RLS, seed script, server-side queries", "Supabase cloud"),
        ],
    )
    add_body(
        doc,
        "Public catalog pages load car data from Supabase on the server. Authenticated users "
        "book vehicles with RLS-scoped inserts. Administrators manage fleet and bookings "
        "through role-enforced middleware and database policies. All Supabase communication "
        "runs on the Next.js server; the browser interacts only with the application domain.",
    )

    add_subheading(doc, "5.4 Scope of the Project")
    add_body(
        doc,
        "DriveEase is a full-stack car rental web application. Data is stored in Supabase "
        "PostgreSQL tables (cars, bookings, profiles) with Row Level Security. User "
        "authentication is handled by Supabase Auth with seeded demo accounts for "
        "development. The application does not integrate payment gateways or email "
        "notifications. It runs locally with npm run dev and can be deployed to Vercel.",
    )
    add_page_break(doc)

    # SYSTEM ANALYSIS
    add_section_heading(doc, "SYSTEM ANALYSIS")
    add_subheading(doc, "6.1 Existing System")
    add_body(
        doc,
        "In a conventional manual car rental setup, customers contact the agency by phone "
        "or in person to inquire about vehicle availability. Staff maintain records in "
        "spreadsheets or paper logs. Pricing is calculated manually. There is no "
        "self-service portal for customers to browse the fleet or view past bookings.",
    )

    add_subheading(doc, "6.2 Proposed System")
    add_body(
        doc,
        "DriveEase replaces the manual workflow with a single web application backed by "
        "Supabase. Visitors browse cars at /cars; logged-in users book at /cars/[id]/book "
        "and view reservations at /my-bookings; administrators manage fleet and bookings "
        "at /admin/*. Data flows through Server Components and Server Actions to Supabase "
        "PostgreSQL. Middleware refreshes the Supabase session and enforces access.",
    )

    add_subheading(doc, "6.3 User Roles")
    add_table(
        doc,
        ["Role", "Description", "Access"],
        [
            ("Visitor", "Unauthenticated user", "Home, car catalog, car detail"),
            ("User", "Registered customer", "All visitor routes plus booking and my bookings"),
            ("Admin", "Fleet manager", "All user routes plus admin dashboard, cars CRUD, bookings"),
        ],
    )

    add_subheading(doc, "6.4 Feasibility Study")
    add_bullet(doc, "Technical feasibility: Next.js 14, React 18, TypeScript, and Supabase provide a mature full-stack ecosystem.")
    add_bullet(doc, "Operational feasibility: Seeded demo credentials allow immediate testing via Supabase Auth.")
    add_bullet(doc, "Economic feasibility: Open-source stack with Supabase free tier for development and deployment.")
    add_bullet(doc, "Operational readiness: Modular separation of presentation, business logic, and Supabase integration supports independent testing and Vercel deployment via environment variables.")

    add_subheading(doc, "6.5 Context Diagram and Data Flow")
    add_body(
        doc,
        "Visitors and users interact with the Next.js application through a web browser. "
        "The Next.js server handles page rendering, Server Actions, and middleware session "
        "refresh. All database and authentication calls go from Next.js to Supabase Cloud. "
        "For catalog reads, Server Components query the cars table and render HTML. For "
        "bookings, Server Actions insert rows with the authenticated user UUID. The browser "
        "Network tab shows requests only to the application domain, not to supabase.co.",
    )
    add_page_break(doc)

    # SYSTEM SPECIFICATION
    add_section_heading(doc, "SYSTEM SPECIFICATION")
    add_subheading(doc, "7.1 Hardware Requirements")
    add_table(
        doc,
        ["Component", "Minimum Requirement"],
        [
            ("Processor", "Dual-core CPU or equivalent"),
            ("RAM", "4 GB (8 GB recommended)"),
            ("Storage", "500 MB free disk space"),
            ("Display", "1280×720 or higher"),
            ("Network", "Internet for npm install and external images"),
        ],
    )

    add_subheading(doc, "7.2 Software Requirements")
    add_table(
        doc,
        ["Software", "Version"],
        [
            ("Node.js", "18+ (20 recommended)"),
            ("Supabase", "Cloud project (PostgreSQL + Auth)"),
            ("npm", "Bundled with Node.js"),
            ("Operating System", "Windows, macOS, or Linux"),
            ("Web Browser", "Chrome, Firefox, Edge, or Safari (latest)"),
        ],
    )

    add_subheading(doc, "7.3 Functional Requirements")
    add_table(
        doc,
        ["Module", "Requirement", "Status"],
        [
            ("Layout", "Navbar, Footer, mobile nav, home hero", "Implemented"),
            ("Catalog", "Car catalog from Supabase, grid, search/filter, detail page", "Implemented"),
            ("Authentication", "Supabase Auth login, session cookies, middleware guards", "Implemented"),
            ("Booking", "Booking form, price calculation, my bookings page", "Implemented"),
            ("Administration", "Admin dashboard, cars CRUD, booking status updates", "Implemented"),
            ("User Experience", "Loading skeletons, 404 page, toasts, empty states", "Implemented"),
            ("Database & Security", "Supabase schema, RLS, seed script, server data layer", "Implemented"),
        ],
    )

    add_subheading(doc, "7.4 Non-Functional Requirements")
    add_bullet(doc, "Responsive design using Tailwind CSS breakpoints.")
    add_bullet(doc, "Server-side data access: all Supabase queries run on the Next.js server.")
    add_bullet(doc, "Server-side validation for booking dates, credentials, and admin mutations.")
    add_bullet(doc, "Semantic HTML and ARIA attributes for accessibility.")
    add_bullet(doc, "Production security: hashed passwords via Supabase Auth, HTTP-only cookies, admin role in JWT app_metadata enforced at middleware and RLS.")
    add_bullet(doc, "Deployment readiness: builds with npm run build; Vercel runtime requires only NEXT_PUBLIC_SUPABASE_URL and NEXT_PUBLIC_SUPABASE_ANON_KEY.")

    add_subheading(doc, "7.5 Constraints and Assumptions")
    add_bullet(doc, "Demo accounts are pre-seeded; no public registration page.")
    add_bullet(doc, "Car images use external URLs; Supabase Storage not used.")
    add_bullet(doc, "Availability checks use car.available flag only, not date overlap.")
    add_bullet(doc, "Assumes Supabase migrations applied and npm run seed executed before testing.")
    add_page_break(doc)

    # SOFTWARE DESCRIPTION
    add_section_heading(doc, "SOFTWARE DESCRIPTION")
    add_subheading(doc, "8.1 Technology Stack")
    add_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ("Framework", "Next.js 14.2 (App Router)", "Routing, SSR, Server Actions"),
            ("Language", "TypeScript 5", "Static typing"),
            ("UI Library", "React 18", "Component-based UI"),
            ("Styling", "Tailwind CSS 3.4", "Utility-first CSS"),
            ("Components", "shadcn/ui", "Button, Card, Table, Sheet, etc."),
            ("Notifications", "Sonner", "Toast messages"),
            ("Backend", "Supabase", "PostgreSQL database + Auth"),
            ("Supabase Client", "@supabase/ssr", "Server-side session and queries"),
            ("Unit Tests", "Vitest 2.1", "Library function tests"),
            ("E2E Tests", "Playwright 1.59", "Browser automation tests"),
            ("Deployment", "Vercel (optional)", "Hosting with Supabase env vars"),
        ],
    )

    add_subheading(doc, "8.2 System Architecture")
    add_body(
        doc,
        "The application follows a layered architecture: (1) Presentation layer (app/ + "
        "components/) for pages and UI; (2) Business logic layer (lib/) for auth and "
        "Supabase data access; (3) Supabase layer (lib/supabase/) for server client and "
        "middleware session refresh; (4) Data layer (Supabase PostgreSQL) with RLS; "
        "(5) Security layer (middleware.ts + RLS) for route and database protection.",
    )
    add_body(
        doc,
        "Server Actions in app/actions/ handle mutations (login, logout, create booking, "
        "admin CRUD) by calling Supabase on the server with the user's session cookie.",
    )
    add_body(
        doc,
        "Three-tier request flow: the browser sends HTTP requests only to Next.js (Vercel or "
        "localhost). Server Components and Server Actions invoke Supabase clients to read or "
        "write PostgreSQL data and manage Auth sessions. Supabase returns results to the server, "
        "which renders HTML or redirects the browser.",
    )

    add_subheading(doc, "8.3 Key Modules")
    add_table(
        doc,
        ["Module", "Path", "Responsibility"],
        [
            ("Auth", "lib/auth.ts", "Map Supabase user to session payload"),
            ("Supabase server", "lib/supabase/server.ts", "Server Supabase client with cookies"),
            ("Supabase middleware", "lib/supabase/middleware.ts", "Refresh auth session"),
            ("Data (read)", "lib/data.ts", "Async car listing from Supabase"),
            ("Car catalog", "lib/car-catalog.ts", "Client-safe filter helpers"),
            ("Cars store (write)", "lib/cars-store.ts", "Admin car CRUD via Supabase"),
            ("Bookings", "lib/bookings.ts", "Read/write bookings with joins"),
            ("Booking math", "lib/booking-math.ts", "Rental days and total price"),
            ("Admin auth", "lib/admin-auth.ts", "Layout-level admin guard"),
            ("Middleware", "middleware.ts", "Session refresh and route protection"),
            ("Seed script", "scripts/seed-supabase.ts", "Seed demo users, cars, bookings"),
        ],
    )

    add_subheading(doc, "8.4 Server-Side vs Client-Side Responsibilities")
    add_body(
        doc,
        "The browser never communicates directly with Supabase. Server Components and "
        "Server Actions fetch and mutate data. Client components handle form UI, in-memory "
        "catalog filtering, and toast notifications only.",
    )
    add_table(
        doc,
        ["Task", "Runs on"],
        [
            ("Load car catalog", "Server"),
            ("Load user bookings", "Server"),
            ("Login and logout", "Server Action"),
            ("Create booking", "Server Action"),
            ("Admin car CRUD", "Server Action"),
            ("Filter cars by fuel/price", "Client (in-memory)"),
        ],
    )

    add_subheading(doc, "8.5 Server Actions Reference")
    add_table(
        doc,
        ["Action", "File", "Effect"],
        [
            ("loginAction", "app/actions/auth.ts", "Supabase sign-in; role redirect"),
            ("logoutAction", "app/actions/auth.ts", "Supabase sign-out"),
            ("createBookingAction", "app/actions/booking.ts", "Insert booking row"),
            ("createCarAction", "app/actions/admin.ts", "Insert car (admin)"),
            ("updateCarAction", "app/actions/admin.ts", "Update car (admin)"),
            ("deleteCarAction", "app/actions/admin.ts", "Delete car (admin)"),
            ("updateBookingStatusAction", "app/actions/admin.ts", "Update status (admin)"),
        ],
    )
    add_page_break(doc)

    # PROJECT DESCRIPTION
    add_section_heading(doc, "PROJECT DESCRIPTION")
    add_subheading(doc, "9.1 Application Routes")
    add_table(
        doc,
        ["Route", "Access", "Description"],
        [
            ("/", "Public", "Home page with hero and Browse Cars button"),
            ("/cars", "Public", "Car catalog with search and filters"),
            ("/cars/[id]", "Public", "Car detail with gallery and Book Now"),
            ("/cars/[id]/book", "Authenticated user", "Booking form with date pickers"),
            ("/login", "Public", "Email/password login for user and admin"),
            ("/my-bookings", "Authenticated user", "Table of user's bookings"),
            ("/admin/dashboard", "Admin", "Statistics and quick actions"),
            ("/admin/cars", "Admin", "Cars list, create, edit, delete"),
            ("/admin/bookings", "Admin", "All bookings with status dropdown"),
        ],
    )

    add_body(
        doc,
        "Public routes (/, /cars, /cars/[id]) load catalog data without authentication. "
        "Auth routes (/login) redirect signed-in users by role. User routes (/cars/[id]/book, "
        "/my-bookings) require middleware session validation. Admin routes (/admin/*) require "
        "app_metadata.role = admin.",
    )

    add_subheading(doc, "9.2 Data Model (Supabase PostgreSQL)")
    add_bullet(doc, "auth.users: managed by Supabase Auth; role in app_metadata (user | admin)")
    add_bullet(doc, "profiles: id (uuid), name, email, created_at")
    add_bullet(doc, "cars: id, name, seats, fuel, price, available, images[], description")
    add_bullet(doc, "bookings: id, user_id, car_id, pickup_date, return_date, total_price, status")
    add_body(doc, "RLS: public car SELECT; users manage own bookings; admins manage cars and all bookings.")

    add_subheading(doc, "9.3 Application Modules and Features")
    modules = [
        "Layout Shell: Root layout with Navbar, Footer, AuthProvider, and Sonner toast provider.",
        "Vehicle Catalog: CarsCatalog with responsive grid, filters, CarCard, and CarImageGallery.",
        "Authentication: /login uses supabase.auth.signInWithPassword; session in HTTP-only cookies via @supabase/ssr.",
        "Booking Workflow: createBookingAction inserts into Supabase bookings; my bookings lists user reservations.",
        "Administration: Dashboard, cars CRUD, and booking status management via Supabase.",
        "User Experience: Loading skeletons, custom 404, toast notifications, and empty states.",
    ]
    for mod in modules:
        add_bullet(doc, mod)

    add_subheading(doc, "9.4 Authentication and Security")
    add_bullet(doc, "Supabase Auth handles email/password sign-in and sign-out.")
    add_bullet(doc, "@supabase/ssr manages session cookies on the Next.js server and in middleware.")
    add_bullet(doc, "Admin role stored in auth.users raw_app_meta_data.role (not user_metadata).")
    add_bullet(doc, "Row Level Security enforces database access per role.")
    add_bullet(doc, "User credentials: user@gmail.com / 123456")
    add_bullet(doc, "Admin credentials: admin@gmail.com / admin123")
    add_body(
        doc,
        "Environment variables: NEXT_PUBLIC_SUPABASE_URL, NEXT_PUBLIC_SUPABASE_ANON_KEY "
        "(required); SUPABASE_SERVICE_ROLE_KEY (local seed script only).",
    )

    add_subheading(doc, "9.5 Booking System")
    add_body(
        doc,
        "Users select pickup and return dates. billableRentalDays calculates inclusive rental "
        "days. computeBookingTotal multiplies days by the car's daily price. createBookingAction "
        "validates dates and availability, inserts into Supabase bookings, and redirects to my bookings.",
    )

    add_subheading(doc, "9.6 Admin Panel")
    add_bullet(doc, "Dashboard: Statistics cards for cars and bookings; latest booking summary.")
    add_bullet(doc, "Cars management: Table listing, add/edit forms, delete with confirmation.")
    add_bullet(doc, "Bookings management: Table with status update dropdown via Server Action.")
    add_body(
        doc,
        "Admin car changes reflect immediately on the public catalog because both read and "
        "write use the same Supabase cars table.",
    )

    add_subheading(doc, "9.7 Supabase Backend Architecture")
    add_body(
        doc,
        "DriveEase persists all application data in Supabase PostgreSQL and authenticates "
        "users through Supabase Auth. Tables profiles, cars, and bookings are linked by "
        "foreign keys and protected by Row Level Security policies.",
    )
    add_bullet(doc, "lib/supabase/server.ts — cookie-aware client for Server Components and Actions.")
    add_bullet(doc, "lib/supabase/middleware.ts — refreshes auth tokens on each request.")
    add_bullet(doc, "lib/supabase/admin.ts — service-role client for npm run seed only.")
    add_bullet(doc, "supabase/migrations/ — initial schema setup SQL for greenfield installation.")
    add_bullet(doc, "scripts/seed-supabase.ts — creates demo users, ten cars, and sample bookings.")
    add_body(
        doc,
        "Environment variables: NEXT_PUBLIC_SUPABASE_URL and NEXT_PUBLIC_SUPABASE_ANON_KEY "
        "(required at runtime); SUPABASE_SERVICE_ROLE_KEY (local seed script only).",
    )

    add_subheading(doc, "9.8 Database Schema and Row Level Security")
    add_bullet(doc, "profiles: users read/update own row; admins read all profiles.")
    add_bullet(doc, "cars: public SELECT for catalog; admin-only INSERT, UPDATE, DELETE.")
    add_bullet(doc, "bookings: users SELECT and INSERT own rows; admins SELECT and UPDATE all.")
    add_body(
        doc,
        "Admin role checks use auth.jwt() app_metadata.role = admin. Profile auto-creation "
        "trigger runs on new auth.users insert for future registration support.",
    )

    add_subheading(doc, "9.9 User Interface Components")
    add_table(
        doc,
        ["Component", "Type", "Responsibility"],
        [
            ("Navbar", "Client", "Navigation, user greeting, logout"),
            ("CarsCatalog", "Client", "Filter controls and car grid"),
            ("CarCard", "Server", "Individual car summary card"),
            ("LoginForm", "Client", "Email/password form with error alert"),
            ("CarBookingForm", "Client", "Date pickers and booking submit"),
            ("AdminNav", "Client", "Admin sidebar navigation"),
        ],
    )

    add_subheading(doc, "9.10 Validation and Error Handling")
    add_bullet(doc, "Invalid login credentials show generic error without revealing which field failed.")
    add_bullet(doc, "Booking rejects invalid dates, unavailable cars, and unauthenticated sessions.")
    add_bullet(doc, "Admin actions validate role, record existence, and allowed status values.")
    add_bullet(doc, "Empty filter results and empty booking lists show dedicated empty states.")
    add_bullet(doc, "Expired session on protected routes redirects to login; non-admin users blocked from /admin/*.")
    add_page_break(doc)

    # SYSTEM TESTING
    add_section_heading(doc, "SYSTEM TESTING")
    add_subheading(doc, "10.1 Unit Testing")
    add_table(
        doc,
        ["Test File", "Coverage"],
        [
            ("lib/data.test.ts", "filterCars, distinct fuels/seats (pure helpers)"),
            ("lib/bookings.test.ts", "billableRentalDays, computeBookingTotal"),
        ],
    )
    add_body(doc, "Run with: npm run test")

    add_subheading(doc, "10.2 End-to-End Testing")
    add_table(
        doc,
        ["Spec File", "Coverage"],
        [
            ("e2e/home.spec.ts", "Hero, CTA, navbar links"),
            ("e2e/cars.spec.ts", "10 cars, search/filter, empty state"),
            ("e2e/car-detail.spec.ts", "Metadata, Book CTA, unavailable car, 404"),
            ("e2e/login.spec.ts", "Form, invalid credentials, user callback, admin redirect"),
            ("e2e/auth-guard.spec.ts", "Guards on my-bookings, admin, book route"),
            ("e2e/bookings.spec.ts", "Seeded bookings list, create booking flow"),
            ("e2e/admin.spec.ts", "Non-admin blocked, dashboard, car CRUD, status update"),
            ("e2e/phase6.spec.ts", "Custom 404, login error feedback"),
        ],
    )
    add_body(doc, "Run npm run seed before e2e tests on a fresh Supabase project. Then: npm run test:e2e")

    add_subheading(doc, "10.3 Manual Test Checklist")
    checklist = [
        "Configure .env.local with Supabase keys; run npm run seed.",
        "Open home page; click Browse Cars; verify 10 cars load from Supabase.",
        "Filter by fuel type and price; verify results update.",
        "Open car detail; click Book Now while logged out; verify redirect to login.",
        "Log in as user@gmail.com; complete a booking; verify my bookings page.",
        "Log out; log in as admin@gmail.com; verify redirect to admin dashboard.",
        "Add, edit, and delete a car in admin cars page.",
        "Update a booking status in admin bookings page.",
        "Visit invalid URL; verify custom 404 page.",
    ]
    for item in checklist:
        add_bullet(doc, item)

    add_subheading(doc, "10.4 Sample Test Cases")
    add_table(
        doc,
        ["TC ID", "Module", "Expected Result"],
        [
            ("TC-01", "Catalog", "10 cars displayed on /cars"),
            ("TC-03", "Auth", "Wrong password shows error alert"),
            ("TC-05", "Auth", "Admin login redirects to dashboard"),
            ("TC-07", "Guard", "User blocked from /admin/cars"),
            ("TC-08", "Booking", "3-day rental computes correct total price"),
            ("TC-11", "Admin", "Deleted car removed from catalog"),
        ],
    )
    add_page_break(doc)

    # SYSTEM IMPLEMENTATION
    add_section_heading(doc, "SYSTEM IMPLEMENTATION")
    add_subheading(doc, "11.1 Installation and Setup")
    add_body(doc, "cd car-rental")
    add_body(doc, "npm install")
    add_body(doc, "cp .env.local.example .env.local  # add Supabase keys")
    add_body(doc, "npm run seed  # after applying supabase/migrations SQL")
    add_body(doc, "npm run dev")
    add_body(doc, "Open http://localhost:3000 in a browser.")

    add_subheading(doc, "11.2 Project Structure")
    structure = """car-rental/
├── app/                    # Next.js App Router pages and layouts
│   ├── actions/            # Server Actions (auth, booking, admin)
│   ├── admin/              # Admin panel routes
│   ├── cars/               # Catalog and detail routes
│   ├── login/              # Login page
│   └── my-bookings/        # User bookings page
├── components/             # React components
├── lib/                    # Business logic and Supabase access
│   └── supabase/           # Server, middleware, admin clients
├── scripts/                # Seed script and seed data
├── supabase/migrations/    # Postgres schema and RLS SQL
├── e2e/                    # Playwright end-to-end tests
├── docs/                   # Project reports
└── middleware.ts           # Session refresh and route protection"""
    add_body(doc, structure)

    add_subheading(doc, "11.3 Build and Deployment")
    add_body(doc, "npm run lint — Run ESLint")
    add_body(doc, "npm run build — Production build")
    add_body(doc, "npm start — Start production server")
    add_body(doc, "Vercel: set NEXT_PUBLIC_SUPABASE_URL and NEXT_PUBLIC_SUPABASE_ANON_KEY; add site URL in Supabase Auth settings.")

    add_subheading(doc, "11.4 Supabase Project Setup")
    setup_steps = [
        "Create Supabase account and new project.",
        "Run SQL migrations from supabase/migrations/ in SQL Editor.",
        "Copy Project URL, anon key, and service role key to .env.local.",
        "Run npm run seed to create demo users, cars, and bookings.",
        "Configure Auth redirect URLs for localhost and Vercel domain.",
    ]
    for step in setup_steps:
        add_bullet(doc, step)

    add_subheading(doc, "11.5 Vercel Deployment Guide")
    add_bullet(doc, "Import GitHub repository into Vercel.")
    add_bullet(doc, "Add NEXT_PUBLIC_SUPABASE_URL and NEXT_PUBLIC_SUPABASE_ANON_KEY env vars.")
    add_bullet(doc, "Do not add SUPABASE_SERVICE_ROLE_KEY to Vercel.")
    add_bullet(doc, "Add Vercel URL to Supabase Auth allowed redirect URLs.")
    add_bullet(doc, "Deploy and verify login, catalog, booking, and admin flows.")

    add_subheading(doc, "11.6 Network Privacy (Browser vs Server)")
    add_body(
        doc,
        "DriveEase ensures the browser does not expose Supabase API traffic. All database "
        "queries and authentication API calls execute on the Next.js server. The browser "
        "Network tab shows requests only to the application domain. Session cookies are "
        "HTTP-only. External requests are primarily Unsplash image URLs for car photos.",
    )

    add_subheading(doc, "11.7 Source Code Repository")
    add_table(
        doc,
        ["Item", "Location"],
        [
            ("Primary repository", "https://github.com/ghariharanit/car-rental"),
            ("Main branch", "main — Supabase-backed application and documentation"),
            (
                "Collaboration branch",
                "https://github.com/sabareesh0609/car-rental/tree/backend-integration",
            ),
            ("Documentation", "docs/driveease-project-report.md, .docx"),
            ("Database setup", "supabase/migrations/"),
            ("Screenshots", "docs/screenshots/ (Appendix G)"),
        ],
    )
    add_body(
        doc,
        "Clone with git clone, configure .env.local from .env.local.example, apply SQL "
        "migrations, run npm run seed, then npm run dev. Refresh report screenshots with "
        "npm run screenshots:report while the dev server is running.",
    )
    add_page_break(doc)

    # CONCLUSION
    add_section_heading(doc, "CONCLUSION AND FUTURE ENHANCEMENT")
    add_subheading(doc, "Conclusion")
    add_body(
        doc,
        "DriveEase successfully implements a full-stack car rental application with a complete "
        "customer journey from browsing the fleet through Supabase-authenticated booking and "
        "viewing reservations, as well as a full admin workflow backed by PostgreSQL.",
    )
    add_body(
        doc,
        "The system uses industry-standard authentication with hashed passwords, database-enforced "
        "Row Level Security, and server-side Supabase data access suitable for cloud deployment "
        "on Vercel. Unit tests validate booking math and catalog filters; end-to-end tests cover "
        "authentication guards, booking flows, and admin operations.",
    )

    add_subheading(doc, "Future Enhancements")
    future = [
        "User registration with email confirmation via Supabase Auth.",
        "Payment gateway integration (Stripe or Razorpay).",
        "Email booking confirmations via Edge Functions.",
        "Date-range availability blocking for overlapping bookings.",
        "Supabase Storage for admin car photo uploads.",
        "OAuth login (Google, GitHub).",
        "Audit logging for admin mutations.",
        "CI/CD pipeline with automated tests on pull requests.",
        "Dashboard analytics charts for booking trends.",
        "Multi-location pickup/drop-off support.",
    ]
    for item in future:
        add_bullet(doc, item)
    add_page_break(doc)

    # APPENDIX
    add_section_heading(doc, "APPENDIX")
    add_subheading(doc, "Appendix A — Route Protection Map")
    add_table(
        doc,
        ["Route Pattern", "Visitor", "User", "Admin"],
        [
            ("/, /cars, /cars/[id]", "Yes", "Yes", "Yes"),
            ("/login", "Yes", "Redirect if logged in", "Redirect if logged in"),
            ("/cars/[id]/book", "Redirect to login", "Yes", "Yes"),
            ("/my-bookings", "Redirect to login", "Yes", "Yes"),
            ("/admin/*", "Redirect to login", "Redirect to home", "Yes"),
        ],
    )

    add_subheading(doc, "Appendix B — Database Schema (Supabase)")
    add_body(doc, "profiles: id uuid PK, name text, email text, created_at timestamptz")
    add_body(doc, "cars: id bigserial PK, name, seats, fuel, price, available, images text[], description")
    add_body(doc, "bookings: id bigserial PK, user_id uuid, car_id bigint, pickup_date, return_date, total_price, status")

    add_subheading(doc, "Appendix C — Row Level Security Policy Matrix")
    add_table(
        doc,
        ["Table", "Operation", "Who"],
        [
            ("profiles", "SELECT", "own user or admin"),
            ("cars", "SELECT", "everyone (anon + auth)"),
            ("cars", "INSERT/UPDATE/DELETE", "admin only"),
            ("bookings", "SELECT/INSERT", "own user"),
            ("bookings", "SELECT/UPDATE", "admin (all rows)"),
        ],
    )

    add_subheading(doc, "Appendix D — Environment Variables")
    add_table(
        doc,
        ["Variable", "Required", "Purpose"],
        [
            ("NEXT_PUBLIC_SUPABASE_URL", "Yes", "Supabase project URL"),
            ("NEXT_PUBLIC_SUPABASE_ANON_KEY", "Yes", "Server-side Supabase client"),
            ("SUPABASE_SERVICE_ROLE_KEY", "Seed only", "Local npm run seed script"),
        ],
    )

    add_subheading(doc, "Appendix E — Server Actions Summary")
    add_table(
        doc,
        ["Action", "Auth", "Supabase operation"],
        [
            ("loginAction", "No", "auth.signInWithPassword"),
            ("logoutAction", "Yes", "auth.signOut"),
            ("createBookingAction", "Yes", "INSERT bookings"),
            ("createCarAction", "Admin", "INSERT cars"),
            ("updateBookingStatusAction", "Admin", "UPDATE bookings"),
        ],
    )

    add_subheading(doc, "Appendix F — Seeded Demo Data")
    add_bullet(doc, "10 cars (IDs 1–10) with mixed fuel types; Ford EcoSport unavailable.")
    add_bullet(doc, "2 users: user@gmail.com (user), admin@gmail.com (admin).")
    add_bullet(doc, "2 sample bookings for demo user.")
    add_bullet(doc, "Fuel breakdown: 4 Petrol, 3 Diesel, 2 Electric, 1 Hybrid.")

    add_subheading(doc, "Appendix G — Application Screenshots")
    add_body(
        doc,
        "UI screenshots from http://localhost:3000 with seeded Supabase demo data. "
        "Regenerate with npm run screenshots:report (requires npm run dev).",
    )
    add_table(
        doc,
        ["Fig.", "Screen", "Route", "File"],
        [
            ("G.1", "Home page", "/", "01-home.png"),
            ("G.2", "Cars catalog", "/cars", "02-cars-catalog.png"),
            ("G.3", "Car detail", "/cars/1", "03-car-detail.png"),
            ("G.4", "Login", "/login", "04-login.png"),
            ("G.5", "My Bookings", "/my-bookings", "05-my-bookings.png"),
            ("G.6", "Booking form", "/cars/1/book", "06-booking-form.png"),
            ("G.7", "Admin dashboard", "/admin/dashboard", "07-admin-dashboard.png"),
            ("G.8", "Admin cars", "/admin/cars", "08-admin-cars.png"),
            ("G.9", "Admin bookings", "/admin/bookings", "09-admin-bookings.png"),
        ],
    )
    figures = [
        ("01-home.png", "Figure G.1 — Home page hero section with Browse Cars button"),
        ("02-cars-catalog.png", "Figure G.2 — Cars catalog with search and filters"),
        ("03-car-detail.png", "Figure G.3 — Car detail with image gallery"),
        ("04-login.png", "Figure G.4 — Login page"),
        ("05-my-bookings.png", "Figure G.5 — My Bookings table"),
        ("06-booking-form.png", "Figure G.6 — Booking form with date pickers"),
        ("07-admin-dashboard.png", "Figure G.7 — Admin dashboard statistics"),
        ("08-admin-cars.png", "Figure G.8 — Admin cars management"),
        ("09-admin-bookings.png", "Figure G.9 — Admin bookings management"),
    ]
    for filename, caption in figures:
        add_figure(doc, filename, caption)
    add_body(
        doc,
        "Figures G.10 (Supabase Table Editor) and G.11 (Vercel deployment) should be "
        "captured manually as 10-supabase-tables.png and 11-vercel-deploy.png.",
    )

    add_subheading(doc, "Appendix H — Glossary")
    glossary = [
        "RLS — Row Level Security; PostgreSQL row filtering per user",
        "Server Component — React component rendered on the Next.js server",
        "Server Action — Server-side mutation invoked from client forms",
        "JWT — JSON Web Token used by Supabase Auth for session claims",
        "UUID — Unique identifier for Supabase auth user IDs",
    ]
    for item in glossary:
        add_bullet(doc, item)

    add_subheading(doc, "Appendix I — npm Scripts Reference")
    add_table(
        doc,
        ["Command", "Description"],
        [
            ("npm run dev", "Start development server"),
            ("npm run build", "Production build"),
            ("npm start", "Start production server"),
            ("npm run lint", "Run ESLint"),
            ("npm run test", "Run Vitest unit tests"),
            ("npm run seed", "Seed Supabase demo users, cars, bookings"),
            ("npm run screenshots:report", "Capture UI screenshots for Appendix G"),
            ("npm run test:e2e", "Run Playwright e2e tests"),
        ],
    )

    add_subheading(doc, "Appendix J — Source Code Listing")
    add_body(
        doc,
        "Main source files for DriveEase. Full repository: "
        "https://github.com/ghariharanit/car-rental (branch main).",
    )
    add_table(
        doc,
        ["Layer", "Path", "Description"],
        [
            ("Routes", "app/page.tsx", "Home page"),
            ("Routes", "app/cars/", "Catalog and car detail pages"),
            ("Routes", "app/admin/", "Admin dashboard, cars, bookings"),
            ("Server Actions", "app/actions/", "Auth, booking, admin mutations"),
            ("Middleware", "middleware.ts", "Session refresh and route guards"),
            ("Supabase", "lib/supabase/", "Server, middleware, admin clients"),
            ("Data", "lib/data.ts, lib/bookings.ts", "Supabase read/write helpers"),
            ("Database", "supabase/migrations/", "Schema and RLS SQL"),
            ("Seed", "scripts/seed-supabase.ts", "Demo users, cars, bookings"),
            ("Tests", "e2e/, lib/*.test.ts", "Playwright and Vitest tests"),
            ("Docs", "docs/, scripts/generate-project-doc.py", "Report and generator"),
        ],
    )
    add_body(doc, "Representative excerpt — middleware.ts:")
    add_body(doc, "export async function middleware(request) { return await updateSession(request); }")
    add_body(doc, "Representative excerpt — createBookingAction validates session before Supabase insert.")
    add_page_break(doc)

    # BIBLIOGRAPHY
    add_section_heading(doc, "BIBLIOGRAPHY & REFERENCES")
    refs = [
        "Next.js Documentation — https://nextjs.org/docs",
        "React Documentation — https://react.dev",
        "TypeScript Documentation — https://www.typescriptlang.org/docs",
        "Supabase Documentation — https://supabase.com/docs",
        "Supabase Auth with Next.js — https://supabase.com/docs/guides/auth/server-side/nextjs",
        "Supabase Row Level Security — https://supabase.com/docs/guides/database/postgres/row-level-security",
        "Tailwind CSS Documentation — https://tailwindcss.com/docs",
        "shadcn/ui — https://ui.shadcn.com",
        "Sonner Toast Library — https://sonner.emilkowal.ski",
        "Vitest — https://vitest.dev",
        "Playwright — https://playwright.dev",
        "Vercel Deployment — https://vercel.com/docs",
        "PostgreSQL Documentation — https://www.postgresql.org/docs/",
        "MDN Web Docs — https://developer.mozilla.org",
    ]
    for i, ref in enumerate(refs, 1):
        add_body(doc, f"{i}. {ref}")

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    try:
        doc.save(str(OUTPUT))
        print(f"Generated: {OUTPUT}")
    except PermissionError:
        fallback = OUTPUT.with_name(f"{OUTPUT.stem}-temp{OUTPUT.suffix}")
        doc.save(str(fallback))
        print(
            f"Could not overwrite {OUTPUT} (file may be open in Word). "
            f"Generated: {fallback}"
        )


if __name__ == "__main__":
    main()
